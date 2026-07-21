#!/usr/bin/env python3
"""Generate Sanlian (三联) topic proposal Word document."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = "/workspace/三联选题提案_2026下半年.docx"


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)


def add_title(doc, text, size=22, bold=True, center=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p.paragraph_format.space_after = Pt(12)
    return p


def add_subtitle(doc, text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "楷体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_after = Pt(18)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 13}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 13))
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_mixed_paragraph(doc, parts, indent=False, first_line=False):
    """parts: list of (text, bold)"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.74)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, parts, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "黑体"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def build():
    doc = Document()
    set_doc_defaults(doc)

    add_title(doc, "给三联的三个选题提案")
    add_subtitle(doc, "——一个每周穿越深圳湾的文学硕士，对2026年下半年的观察")

    add_heading(doc, "三个我想做的选题", 1)

    # Topic 1
    add_heading(doc, "选题一：《周末去香港：一亿人次背后的「48小时离岸实验」》", 2)
    add_mixed_paragraph(doc, [
        ("2026年上半年，经深圳各口岸出入境人员突破", False),
        ("1.4亿人次", True),
        ("，单日客流峰值14次破百万；同期访港旅客约", False),
        ("2671万人次", True),
        ("，其中内地旅客占比", False),
        ("77%", True),
        ("。但这一轮赴港潮和十年前完全不同——Mastercard数据显示，内地旅客的购物开支占比已从2018年的", False),
        ("70.5%", True),
        ("骤降至2026年一季度的", False),
        ("44%", True),
        ("。年轻人不去海港城，去深水埗；不逛DFS，逛独立书店；不买金饰，买一杯手冲和一张渡轮船票。", False),
        ("「周末去香港」", True),
        ("正在从消费行为变成一种生活方式的试穿：不辞职、不移民，只是每周给自己48小时「不属于日常」的时间。这背后是一代人对「确定性生活」的疲倦，和对「附近」的重新发现。", False),
    ], first_line=True)

    # Topic 2
    add_heading(doc, "选题二：《新中式走到了岔路口》", 2)
    add_mixed_paragraph(doc, [
        ("2026年，国风服饰市场规模已突破", False),
        ("3000亿元", True),
        ("，新中式服饰规模突破", False),
        ("2500亿元", True),
        ("，全国非遗相关市场规模在2025年已破", False),
        ("1.1万亿", True),
        ("。观夏把东方香气做成了排队经济，霸王茶姬用国风纹样铺满了购物中心，北京时装周上非遗刺绣与工装版型正在碰撞出新的火花。但", False),
        ("审美疲劳已经到来", True),
        ("——当每一家奶茶店都有书法logo，每一个商场都有枯山水，「拼接传统元素的粗放式设计已被市场淘汰」（汉服品牌织造司创始人语），新中式正在从美学变成套路。下一步是什么？一些更安静的尝试正在发生：把宋代点茶变成日常饮品的茶室、用活字印刷做婚礼请柬的工作室、在景德镇烧自用碗碟的年轻陶艺师。", False),
        ("从央视非遗晚会上的云锦高定，到中国国际时装周南京站四大名锦首次同台——", True),
        ("新中式的2.0，也许不再是「看起来很中国」，而是「用起来很日常」。", True),
    ], first_line=True)

    # Topic 3
    add_heading(doc, "选题三：《年轻人开始「求签」了》", 2)
    add_mixed_paragraph(doc, [
        ("据研究机构测算，至2026年，我国寺庙经济市场规模有望突破", False),
        ("千亿元", True),
        ("大关，年增长率约", False),
        ("10%", True),
        ("。雍和宫的香灰手串在闲鱼溢价十倍，灵隐寺的十八籽手串曾被炒至三倍价，寺庙咖啡馆在社交平台获得的关注度超过任何第三波精品咖啡店。近三年寺庙18-35岁游客占比达", False),
        ("63%", True),
        ("，门票订单量同比涨超", False),
        ("310%", True),
        ("；2026年寺庙义工与短期禅修的搜索量同比暴涨", False),
        ("327%", True),
        ("，85%的报名者是20-30岁的年轻人。有意思的是，这批涌入寺庙的年轻人并不信佛——据调研，", False),
        ("72%", True),
        ("的年轻人去寺庙的核心目的是「缓解焦虑、平复心情」。当考公、考研、求职的确定性路径越来越拥挤，「求个心安」本身变成了一种理性消费。这不是迷信的回潮，是", False),
        ("「精神急诊室」", True),
        ("式的时代症候——确定性稀缺时代里，一种低成本的心理对冲。", False),
    ], first_line=True)

    doc.add_page_break()

    # Full expansion Topic 1
    add_title(doc, "选题一展开", 18, True, True)
    add_subtitle(doc, "《周末去香港：一亿人次背后的「48小时离岸实验」》")

    add_heading(doc, "一、为什么是现在", 1)
    add_mixed_paragraph(doc, [("几组来自2026年的数字，勾勒出一个清晰的趋势：", False)], first_line=True)

    add_mixed_paragraph(doc, [
        ("据深圳市口岸办数据，2026年第一季度深圳口岸累计出入境人员达", False),
        ("7263.3万人次", True),
        ("，同比增长", False),
        ("14.2%", True),
        ("。截至5月2日，2026年经深圳各口岸出入境人员突破", False),
        ("1亿人次", True),
        ("，比2025年提前", False),
        ("17天", True),
        ("达成，日均客流", False),
        ("82.5万人次", True),
        ("——稳居全国第一。截至6月底，这一数字已超过", False),
        ("1.4亿人次", True),
        ("。上半年，单日客流峰值14次突破百万，最高达", False),
        ("109.2万人次", True),
        ("。皇岗、福田两口岸上半年查验出入境人员4100万人次；罗湖口岸日均客流突破20万。深圳湾口岸在周末高峰时段的排队时间一度超过", False),
        ("两小时", True),
        ("，通关大厅的密度堪比春运。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("与此同时，香港旅发局数据显示，2026年上半年访港旅客共录得约", False),
        ("2671万人次", True),
        ("，按年增长", False),
        ("13%", True),
        ("。其中中国内地旅客", False),
        ("2056万人次", True),
        ("，同比显著增长", False),
        ("16%", True),
        ("，占总旅客量约", False),
        ("77%", True),
        ("。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("但真正值得注意的，不是「来的人变多了」，而是", False),
        ("「来的人变了」", True),
        ("。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("Mastercard经济研究所数据显示，2026年第一季度，内地旅客的购物开支占总旅游消费的比例，已由2018年的", False),
        ("70.5%", True),
        ("大幅下降至", False),
        ("44%", True),
        ("——从「纯购物」转向了餐饮、文化和体验式消费。过夜旅客平均消费近", False),
        ("9400元", True),
        ("人民币，即日来回旅客平均消费约", False),
        ("2900元", True),
        ("，而大湾区当日往返的旅客人均花费约", False),
        ("1000元", True),
        ("。翻译成日常语言就是——来的人更多了，停留时间更短了，", False),
        ("但他们不买包了", True),
        ("。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [("打开小红书搜索「香港周末」，你会看到一种全新的叙事：", False)], first_line=True)
    add_mixed_paragraph(doc, [
        ("不是太平山顶的游客照，而是西营盘一家只有六个座位的咖啡馆。不是维多利亚港的夜景，而是深水埗鸭寮街的旧货摊。不是「香港购物攻略」，而是「一个人在长洲发呆的一天」。2026年「五一」黄金周期间，约", False),
        ("98万", True),
        ("内地旅客经各口岸进入香港，较去年增长约", False),
        ("7%", True),
        ("——但不同的是，越来越多人绕过了铜锣湾和尖沙咀，直接前往长洲、南丫、大澳这些离岛。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("这些年轻人正在", False),
        ("把香港重新发明", True),
        ("。他们发明的那个香港，和任何一个旅游指南上的香港都不一样。", False),
    ], first_line=True)

    add_heading(doc, "二、核心洞察：不是旅游，是「离岸」", 1)
    add_mixed_paragraph(doc, [("仔细观察这批「周末港漂」的行为模式，会发现一个有趣的悖论：", False)], first_line=True)

    add_mixed_paragraph(doc, [
        ("他们并不是在「旅行」。", True),
        ("旅行意味着计划、攻略、打卡、朋友圈九宫格。但他们中的很多人，去香港的方式更接近「下楼散步」——没有明确目的地，不做详细攻略，甚至不发朋友圈。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("他们也不是在「消费」。", True),
        ("一个典型的周末港漂，花费可能不到500元人民币：一张高铁票（75元），一碗牛腩面（45港币），一杯冰美式（40港币），一张去长洲的渡轮船票（15港币），一晚青旅或深夜返程。这个消费画像，恰好和Mastercard数据里「大湾区当日往返旅客人均1000元」的现象完全吻合。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("他们在做的事情，更接近一种「临时下线」。", True),
        ("从自己的生活里登出48小时——用一种不同的货币、走一条不同的街道、听一种混杂了粤语和英语的声音、在一个没有健康码和刷脸支付的环境里，短暂地变成「另一个版本的自己」。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("我曾和一位每月至少去香港两次的深圳互联网从业者聊过。她说了一句很准确的话：「不是香港有什么特别好的东西。是在那边的时候，我不需要做自己。」", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("这句话指向了一个更深层的需求——不是对香港的向往，而是对日常的倦怠。当工作日被「效率」填满，当一线城市的每一个角落都被算法触达，「周末去香港」提供的核心产品其实不是香港本身，而是", False),
        ("「一段不被自己的生活追踪到的时间」", True),
        ("。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("我把它叫做", False),
        ("「离岸生活」", True),
        ("：不是移民那种彻底的离开，而是一种周期性的、低成本的、可逆的「暂时不在这里」。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("一个能佐证这一趋势的细节是：2026年「五一」期间，中环码头搭乘渡轮前往长洲的乘客为数众多，慢船接近满座。有长洲度假屋职员表示，假期期间的房间在四月中就已被订满，其中大部分是内地旅客。", False),
        ("「周末去离岛」", True),
        ("，正从一个小众选项变成主流选择。", False),
    ], first_line=True)

    add_heading(doc, "三、一座没有汽车的岛：离岛为什么成了终极目的地", 1)
    add_mixed_paragraph(doc, [
        ("如果说「周末去香港」是第一层离岸，那么", False),
        ("「去香港的离岛」就是离岸的平方", True),
        ("。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("2026年七月的一个周末，我坐快船去长洲。从中环五号码头出发，三十分钟——这是长洲和中环之间的距离，比大部分香港人的通勤时间还短。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("但到达之后的世界完全不同。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("长洲没有汽车。", True),
        ("这句话需要被强调，因为它意味着的东西比字面多得多。没有汽车意味着没有红绿灯、没有马路牙子、没有导航播报的「前方五百米右转」。岛上的交通工具只有三种：自行车、消防车、你的脚。巷子窄到两个人迎面走过需要侧身，头顶是低矮的骑楼和晾衣绳，脚下是被雨水冲刷得发亮的石板。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("这种物理空间的尺度变化，会直接改变人的心理状态。你不再「赶路」，因为岛上没有任何地方需要「赶」着去。你开始注意到码头边渔船上的锈迹、玉虚宫门前石阶上晒太阳的两只土狗、饼店门口堆成小山的平安包上面那个红色的印。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("长洲的平安包是一个绝佳的隐喻。", True),
        ("每年农历四月初八的太平清醮，是香港唯一保留至今的大型民间醮会。全岛吃素三天，最后一晚在北帝庙前竖起包山，年轻人攀上去抢包——抢得越高，福气越大。但在一年中其余的三百六十四天里，平安包只是饼店里的一种普通点心，安静地堆在玻璃柜里，等下一个游客把它带走。", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("一个古老的仪式，一种日常的食物，在同一个容器里并存。这恰好是长洲——以及所有香港离岛——吸引内地年轻人的真正原因：它们保存着一种", False),
        ("「未被优化的生活」", True),
        ("。没有被算法挑选、没有被效率重塑、没有被连锁品牌覆盖。它们只是在那里，按照自己的节奏运转，像一座城市的时间胶囊。", False),
    ], first_line=True)

    add_heading(doc, "四、内容框架", 1)
    add_mixed_paragraph(doc, [
        ("如果做成一期三联专题，建议以", False),
        ("「从近到远的五重离岸」", True),
        ("为结构：", False),
    ], first_line=True)

    add_mixed_paragraph(doc, [
        ("主稿 | ", True),
        ("《48小时离岸：周末去香港的年轻人在逃往哪里》", True),
    ])
    add_mixed_paragraph(doc, [
        ("从深圳湾口岸周末早晨的人潮写起，画出「周末港漂」的群体画像与心理动机。结合2026年上半年1.4亿人次通关、内地购物占比骤降至44%等关键数据，核心访谈5-8位典型人物：互联网打工人、自由职业者、考研二战学生、带孩子的新手妈妈——他们去的是同一个香港，但「离岸」的理由各不相同。", False),
    ], indent=True)

    add_mixed_paragraph(doc, [
        ("分稿一 | ", True),
        ("《重新发明的香港：深水埗、西营盘与一种「反攻略」式旅行》", True),
    ])
    add_mixed_paragraph(doc, [
        ("内地年轻人正在绕过所有旅游攻略推荐的地方，自己「发现」一个新香港。深水埗的布料市场、西营盘的海味街、油麻地的旧书店、北角的福建菜馆——这些本地人习以为常的日常，为什么变成了内地游客眼中的「宝藏」？采访这些街区的老店经营者，和频繁光顾的内地访客。", False),
    ], indent=True)

    add_mixed_paragraph(doc, [
        ("分稿二 | ", True),
        ("《没有汽车的岛：香港离岛的「慢时间」经济学》", True),
    ])
    add_mixed_paragraph(doc, [
        ("长洲、南丫岛、大澳、坪洲——这些曾被视为「偏远」的离岛，正在成为都市年轻人的精神飞地。2026年「五一」期间长洲度假屋提前半个月订满的现象，只是这股潮流的一个切片。但「慢」是有代价的：离岛的年轻人在外流，老人在留守，旅游带来的是复兴还是另一种消耗？深入一座岛屿的日常生活，看看「慢时间」在被消费的同时，是否也在被消解。", False),
    ], indent=True)

    add_mixed_paragraph(doc, [
        ("分稿三 | ", True),
        ("《「松弛感」是一种进口商品吗》", True),
    ])
    add_mixed_paragraph(doc, [
        ("文化评论。小红书上的「港式松弛感」几乎成了一种固定滤镜：旧式冰室的塑料凳、叮叮车窗外的霓虹、维港渡轮上吹风的侧脸。但「松弛」真的是香港的底色吗？对于本地人来说，香港是全球工时最长的城市之一，是六十平米住四口人的逼仄，是没有退休保障的焦虑。内地年轻人感受到的「松弛」，可能不是香港本来的样子，而是「不是我的生活」这件事本身带来的松弛。对话文化学者，拆解这层投射。", False),
    ], indent=True)

    add_heading(doc, "五、采访对象方向", 1)
    add_table(doc, ["类型", "方向"], [
        ["学者", "香港中文大学传播与文化研究学者 / 城市研究学者 / 港大社会学系"],
        ["在地观察者", "香港离岛民宿经营者 / 深水埗老店第二代 / 独立书店主理人（如见山书店、序言书室）"],
        ["内地访客", "5-8位「月均赴港1次以上」的25-35岁人群，覆盖不同职业和动机"],
        ["新媒体", "2-3位小红书/B站「香港生活」头部博主"],
        ["品牌/商业", "在港开店的内地新消费品牌（如% Arabica、话梅等）/ 香港本地新兴文化空间"],
    ])

    add_heading(doc, "六、传播设想", 1)
    add_mixed_paragraph(doc, [("纸刊+公众号：", True), ("封面专题主发布", False)])
    add_mixed_paragraph(doc, [("短视频矩阵（抖音/视频号/B站）：", True)], first_line=True)
    add_bullet(doc, [("「48小时香港」系列vlog，3-5支，跟拍不同人的周末赴港路线", False)])
    add_bullet(doc, [("一支离岛纪录短片（长洲/南丫/大澳任选一岛），5-8分钟", False)])
    add_mixed_paragraph(doc, [("小红书：", True)], first_line=True)
    add_bullet(doc, [("发起 #我的离岸周末 话题，联动KOL做UGC征集", False)])
    add_bullet(doc, [("配合专题发布一组「周末去香港的100个理由」图文", False)])
    add_mixed_paragraph(doc, [("播客：", True)], first_line=True)
    add_bullet(doc, [("可合作「随机波动」或「不合时宜」做一期深聊：为什么我们需要「暂时不在这里」", False)])
    add_mixed_paragraph(doc, [("线下活动：", True)], first_line=True)
    add_bullet(doc, [("深圳/香港双城沙龙，邀请受访者与读者面对面", False)])
    add_mixed_paragraph(doc, [("品牌合作可能性：", True)], first_line=True)
    add_bullet(doc, [("文旅类：香港旅发局、广深港高铁", False)])
    add_bullet(doc, [("消费品类：观夏（离岛叙事 × 东方香气的高度匹配）、三顿半（「周末即刻」概念）", False)])
    add_bullet(doc, [("出行/金融类：AlipayHK、大湾区跨境金融产品", False)])

    add_heading(doc, "七、为什么这个选题适合三联", 1)
    items = [
        ("它有当代性。", "一亿人次的通关数据、内地购物占比骤降26个百分点的消费转向、「五一」98万人涌港的黄金周现象——「周末去香港」是2026年最活跃的青年迁徙现象之一，自带话题热度。"),
        ("它有人文性。", "不停留在攻略和打卡，追问「我们为什么需要暂时离开自己的生活」——这是三联一直在做的事。"),
        ("它有结构感。", "「五重离岸」的框架（口岸→旧城区→离岛→文化想象→自我投射），从具体到抽象，层层推进。"),
        ("它有商业化空间。", "文旅、消费、出行三大品类都有天然的结合点，可以支撑一个完整的商业化项目。"),
        ("它有延展性。", "「离岸生活」可以做成一个持续IP：周末去香港、周末去大理、周末去松阳——「不辞职也能过另一种生活」，这个母题足够支撑一个年度系列。"),
    ]
    for label, body in items:
        add_mixed_paragraph(doc, [(label, True), (body, False)], first_line=True)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
