#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate Heytea × 《给阿嬷的情书》 commercial proposal (v5 commercial goals)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT = Path("/workspace/商业策划案_心中念你_喜茶×给阿嬷的情书.docx")
MOCKUP_DIR = Path("/workspace/proposal-mockups")


def set_run_font(run, size=11, bold=False, color=None, name="宋体"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 18, True, (26, 26, 26), "黑体")
    p.paragraph_format.space_after = Pt(4)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 11, False, (90, 90, 90))
    p.paragraph_format.space_after = Pt(8)


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, 14, True, (26, 26, 26), "黑体")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, 12, True, (40, 40, 40), "黑体")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def body(doc, text, first_indent=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, 11)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(4)
    if first_indent:
        pf.first_line_indent = Cm(0.74)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    r = p.add_run(text)
    set_run_font(r, 11)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def quote(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 12, True, (70, 70, 70), "楷体")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, 9, False, (110, 110, 110))
    p.paragraph_format.space_after = Pt(6)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, 10, True, name="黑体")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, 10)
    doc.add_paragraph()


def try_add_image(doc, path: Path, width_inches=5.8, caption=None):
    if not path.exists():
        note(doc, f"（附图缺失：{path.name}）")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    p.paragraph_format.space_after = Pt(4)
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        set_run_font(r, 9, False, (110, 110, 110))


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.4)
        s.right_margin = Cm(2.4)

    # Cover
    add_title(doc, "喜茶 × 《给阿嬷的情书》")
    add_title(doc, "联名商业策划案")
    add_subtitle(doc, "Campaign：《心中念你》")
    add_subtitle(doc, "Hero Idea：给想念的人，一杯茶")
    add_subtitle(
        doc,
        "版本：商业目标增强版 v5｜面试作业｜仅供内部讨论｜" + date.today().isoformat(),
    )
    quote(doc, "心中念你，便不觉远。")
    body(
        doc,
        "项目定位：从电影「侨批文化」转化为年轻消费者日常表达思念、并把想念送出去的消费场景。",
        False,
    )
    body(
        doc,
        "完整链路：文化资产 → 购买场景 → 情绪表达 → 社交传播 → 会员/内容资产",
        False,
    )
    note(
        doc,
        "说明：本方案为面试作业，非已签约合作。凡涉及票房、联名策略等公开信息均标注来源；价格/KPI为目标体系与模型示意，非喜茶内部真实数据，正式立项需回填。",
    )

    # ========== 01 ==========
    h1(doc, "01｜项目判断：为什么值得做")
    body(doc, "联名市场正从「流量联名」进入「资产联名」：过去是 IP知名度→短期销量；现在是 文化符号→品牌记忆→长期资产。")
    body(
        doc,
        "喜茶已有年轻化、创新茶饮与东方美学基础；用户情绪认知却更多停留在「好喝、好看、年轻」。本片可补齐「有温度、有牵挂、有人情味」。",
    )
    body(
        doc,
        "公开数据支撑窗口存在：据36氪转引猫眼专业版，截至2026-06-27影片上映59天票房破19亿元；据新浪转引猫眼，截至2026-06-30影片累计票房约19.27亿元、猫眼购票评分9.7，为2026上半年购票评分冠军之一；猫眼专业版实时页显示上映约89天累计票房约19.98亿元（长尾仍在累积）。同时，阿嬷手作、一杯潮茶等已出现真实联名先例，而喜茶尚未与该片官联。",
    )
    note(
        doc,
        "数据来源：36氪快讯《电影〈给阿嬷的情书〉总票房破19亿元》（转引猫眼，2026-06-27）；新浪新闻转引猫眼专业版（截至2026-06-30）；猫眼专业版实时票房页（查阅时点约2026-07下旬，上线前请复核最新值）。联名先例来源：阿嬷手作公众号、一杯潮茶抖音（需复核）。",
    )
    body(
        doc,
        "决策一句话：这不是卖一杯奶茶，而是建立「喜茶 = 年轻人表达情感的入口」；文化如何转化商业，是本案核心命题。",
    )

    # ========== 02 ==========
    h1(doc, "02｜项目目标：做完之后得到什么")
    h2(doc, "商业目标")
    bullet(doc, "提升联名新品销售（核心款「念你」承担销量）。")
    bullet(doc, "提升节日节点消费场景（母亲节/中秋/春节等可复用）。")
    bullet(doc, "拉动礼赠与关系消费（从自饮升级为「送给TA」）。")
    bullet(doc, "增加年轻用户会员沉淀（电子侨批/纪念卡绑定会员与二次触达）。")
    h2(doc, "品牌目标")
    bullet(doc, "强化喜茶东方文化表达（东方文化 × 年轻表达）。")
    bullet(doc, "建立情绪价值认知：会念你、有温度、有人情味。")
    h2(doc, "用户目标")
    bullet(doc, "给年轻消费者提供低压力表达思念的方式。")
    bullet(doc, "提供可送达的情绪载体：一杯茶 = 一封可寄出的「侨批」。")
    body(
        doc,
        "成功标准（决策用）：销售（联名销量/套餐比/礼赠客单）+ 用户（参与量/新会员/二次触达）+ 内容（UGC/话题）+ 品牌（情绪认知提升，可用前后测/社媒语义观察）。",
    )

    # ========== 03 ==========
    h1(doc, "03｜用户洞察：谁会买，为什么买")
    h2(doc, "核心用户")
    body(
        doc,
        "25–35岁城市青年（学生尾段/新白领为主）。共同状态：离开家乡进入城市生活；与父母/长辈关系亲密，但表达不足——不是没有爱，而是不知道如何开口。",
    )
    body(
        doc,
        "「念你」提供的是低压力表达入口：买的不是奶茶本身，而是一次可以说出口、送得出去的表达。",
    )
    note(
        doc,
        "人群参考：公开报道称喜茶核心用户偏18–30岁年轻群体（广告门对喜茶×星星人联名分析，2025-12）。本案将核心表达人群收束为25–35，更贴近「异地亲情/礼赠」决策者；正式立项可用喜茶会员画像回填。",
    )
    h2(doc, "洞察→策略")
    table(
        doc,
        ["洞察", "表现", "策略"],
        [
            ["想表达，怕沉重", "不愿公开煽情", "侨批体短笺 / 金句模板 / 轻提示句"],
            ["想送礼，缺入口", "异地关系不知送什么", "「给想念的人，一杯茶」赠送链路"],
            ["要理由才二购", "首次尝鲜后缺复购钩子", "节点礼赠 + 会员权益二次触达"],
            ["分享要真也好看", "拒绝尴尬打卡任务", "电子侨批一键生成与分享"],
        ],
    )

    # ========== 04 ==========
    h1(doc, "04｜市场与品牌机会")
    body(doc, "情绪消费阶段：联名购买动机含情绪价值、身份表达、社交分享。竞争关键不是更热闹，而是更轻、更真、更能日常复用。")
    h2(doc, "喜茶联名策略机会：喜欢 vs 想念")
    table(
        doc,
        ["类型", "代表", "优势", "本案差异"],
        [
            ["强识别IP", "Chiikawa、星星人", "快速吸引年轻用户，制造「喜欢」", "偏萌系陪伴与即时热度"],
            ["文化IP", "《给阿嬷的情书》", "建立情绪资产，制造「想念」", "补齐有温度的东方表达"],
        ],
    )
    body(
        doc,
        "背景判断：2025年喜茶联名呈现「降频提质」——公开报道指其联名从高频转向更少但更重品牌契合的合作（如深圳新闻网2025-12-25；广告门称2025年重点IP联名包括Chiikawa与星星人）。本案契合「从短期流量到品牌心智/情绪资产」的方向，而非再堆一次萌系流量。",
    )
    note(
        doc,
        "来源：深圳新闻网《喜茶联名“降频提质”…》（2025-12-25）；广告门《从chiikawa到星星人…》（2025-12）。联名次数等媒体表述不一，引用时注明「公开报道」，上线前复核。",
    )

    # ========== 05 ==========
    h1(doc, "05｜IP资产矩阵：电影能给喜茶什么")
    table(
        doc,
        ["IP资产", "情绪价值", "商业转化"],
        [
            ["侨批", "跨越距离的思念", "包装/杯套、电子侨批、赠送链路"],
            ["家书", "未表达的情感", "UGC、信箱书写、侨批翻译"],
            ["潮汕茶文化", "东方生活方式", "产品叙事、区域试点记忆点"],
            ["等待与归来", "代际关系与报平安", "节点战役、品牌内容资产"],
        ],
    )
    body(
        doc,
        "授权金句池（正式授权后使用）：「江海万里，心中念你，便不觉遥远。」「暹罗虽远，心有所寄，身若比邻。」「纸短情长，伏惟珍重。」「切要平安，即为团圆。」「江海有岸，团圆可盼。」「展信安康。」",
    )
    body(doc, "禁用边界：不消费丧亲创伤；不夸大真实侨批史；不把「阿嬷」做成低俗怀旧梗。")

    # ========== 06 ==========
    h1(doc, "06｜合作价值")
    table(
        doc,
        ["对象", "价值"],
        [
            ["喜茶·销售", "联名销量、礼赠客单、节点场景扩容"],
            ["喜茶·用户", "会员沉淀与二次触达"],
            ["喜茶·品牌", "东方情绪表达资产与可复用方法论"],
            ["电影IP", "影院后生命周期、年轻非影迷触达、侨批文化大众化"],
        ],
    )
    body(doc, "合作关系：喜茶执行与渠道主体；电影官方授权与审核；不引入第三方中介。")

    # ========== 07 ==========
    h1(doc, "07｜Campaign架构：《心中念你》")
    quote(doc, "Campaign名：《心中念你》｜消费者行动（Hero Idea）：给想念的人，一杯茶")
    body(
        doc,
        "层级关系：Campaign提供品牌命题；Hero Idea提供消费者行动。把一杯茶，变成一封可以被送达的侨批。",
    )
    table(
        doc,
        ["核心动作", "空间", "关键词", "解决什么"],
        [
            ["① 给想念的人，一杯茶", "线上（主传播）", "送出去", "距离与关系消费"],
            ["② 念你信箱", "线下（体验）", "写下来", "表达与停留"],
            ["③ 许多句念你", "内容（资产）", "留下来", "长期内容沉淀"],
        ],
    )
    quote(doc, "一个帮你把想念送出去，一个帮你把想念留下来。")
    body(doc, "链路：购买场景 → 情绪表达 → 社交传播 → 品牌/用户资产。")

    # ========== 08 Hero ==========
    h1(doc, "08｜核心动作1｜给想念的人，一杯茶（Hero Idea）")
    h2(doc, "为什么做")
    body(
        doc,
        "电影核心不是「写信」本身，而是：距离存在时，人如何表达牵挂。侨批解决过去「人在海外，如何把思念寄回家」；喜茶解决现在「人在城市，如何把一份想念送到另一个人手中」。逻辑一致。",
    )
    body(
        doc,
        "年轻人不是没有想念，而是不知道如何表达；异地关系中「想送点什么」缺少低门槛入口。一杯茶 = 可被送达的情绪载体。",
    )
    h2(doc, "体验原则")
    body(doc, "不另造复杂独立App；嵌入喜茶已有APP/小程序消费路径。")
    h2(doc, "同城赠送流程")
    bullet(doc, "打开喜茶APP/小程序 → 选择「送给TA」")
    bullet(doc, "选择关系：家人 / 朋友 / 爱人 / 想念的人")
    bullet(doc, "选择饮品：「念你」")
    bullet(doc, "填写一封「侨批」（三种模式）→ 下单送达")
    h2(doc, "侨批三种写法（降低表达门槛）")
    table(
        doc,
        ["模式", "用户做什么", "系统产出"],
        [
            ["侨批翻译（推荐包装）", "写现代口语思念", "转换为「侨批体」电子信件（勿主打「AI写信」噱头）"],
            ["用户手写", "直接输入原话", "套入侨批视觉版式"],
            ["电影金句模板", "选授权金句+署名", "标准化轻表达"],
        ],
    )
    body(
        doc,
        "侨批翻译示例：用户写「想给妈妈说最近工作很忙但很想她」→ 生成「妈妈展信安康。最近虽忙于城市奔波，心中一直惦念家中三餐四季。愿您身体康健，平安喜乐。」",
    )
    h2(doc, "异地送达：三种商业现实模式")
    table(
        doc,
        ["模式", "适用", "流程", "商业逻辑"],
        [
            ["① 外卖跨店配送（核心）", "双方城市均有喜茶", "A地下单→填B地地址→就近门店制作→骑手配送", "复用已有外卖体系"],
            ["② 电子侨批+自取兑换", "对方附近暂不便配送", "购「念你电子礼」→发电子侨批→对方门店兑换", "类似礼品卡/券逻辑"],
            ["③ 线下同行双杯", "朋友/情侣同行", "「两个人的念你」套餐+双人纪念卡", "提高连带与到店氛围"],
        ],
    )
    body(doc, "覆盖：异地家人、异地朋友、线下关系——不假设「全国随便送」。")
    h2(doc, "商业价值")
    bullet(doc, "拉动赠送消费：自饮 → 关系消费，提高客单。")
    bullet(doc, "创造新场景：从「我想喝」到「我想给某人送一杯」。")
    bullet(doc, "建立用户/内容资产：电子侨批可沉淀、可二次触达。")

    # ========== 09 产品 ==========
    h1(doc, "09｜产品策略与销售场景")
    h2(doc, "产品目标")
    body(doc, "不是只追爆款新品，而是创造「一杯具有购买理由的情绪产品」。名：「念你」。定位：承载思念的东方牛乳茶。")
    h2(doc, "购买三层路径")
    table(
        doc,
        ["层", "消费者获得", "对应"],
        [
            ["口味吸引", "想喝、回甘", "焙香红茶牛乳：熟悉、安心、回甘"],
            ["故事理解", "懂为什么存在", "侨批需要载体；茶亦是连接媒介"],
            ["情绪表达", "可送、可说出口", "送给想念的人；报平安"],
        ],
    )
    body(doc, "桥：包装/杯套完成口味→故事；电子侨批/短笺完成故事→表达。不强调复杂原料创新，优先可量产。")
    h2(doc, "产品矩阵与示意价格带")
    table(
        doc,
        ["组合", "产品", "目的", "示意价格带*"],
        [
            ["单杯", "「念你」", "降低门槛、承担销量", "常规中杯价位带"],
            ["双杯", "「两个人的念你」", "情侣/朋友/亲子", "略低于两杯单买"],
            ["季节限定", "「念你安康」", "节点传播与复购", "限定溢价或同价换装"],
            ["礼赠", "「一封念你礼」", "提高客单", "券+家书卡+明信片，明显高于单杯"],
        ],
    )
    note(doc, "*价格为模型示意，非内部价；以成本、城市价位带与竞品对标后确定。")
    h2(doc, "消费场景与复购设计（为什么买第二次）")
    table(
        doc,
        ["场景", "触发", "购买", "复购钩子"],
        [
            ["日常", "突然想起家里", "单杯", "电子纪念卡收藏 / 轻提示复访"],
            ["节日", "母亲节、中秋、春节", "双杯/礼盒/念你安康", "节点资产复用，不依赖电影档期"],
            ["异地关系", "朋友、情侣、家人", "送给TA / 电子礼", "关系消费场景可反复发生"],
            ["送礼", "给父母/长辈", "一封念你礼", "礼赠理由明确，客单更高"],
        ],
    )
    bullet(doc, "票根核销：试点城市先测核销率与门店成本，再扩面。")

    # ========== 10 门店 ==========
    h1(doc, "10｜核心动作2｜念你信箱（线下增长工具）")
    body(doc, "不是打卡装置，而是门店增长工具。三目标：提升停留时间；提升分享率；提升复购理由。")
    h2(doc, "会员资产链路（经营价值）")
    body(
        doc,
        "写信 → 扫码生成电子侨批/纪念卡 → 绑定喜茶会员 → 获得下一次消费权益（如联名期优惠券/积分示意）→ 二次触达。",
        False,
    )
    body(doc, "路径：购买 → 情绪参与 → 会员沉淀 → 社交传播 → 复购。")
    bullet(doc, "短笺：空白句 + 轻提示句；员工话术一句；不做复杂任务。")
    bullet(doc, "与线上关系：线上「送出去」，线下「写下来」——双空间，不冲突。")

    # ========== 11 内容 ==========
    h1(doc, "11｜核心动作3｜内容传播与资产沉淀")
    bullet(doc, "主话题：#心中念你；行动话题：#给想念的人一杯茶")
    bullet(doc, "辅助：#喝杯茶念一念家里的人")
    h2(doc, "「许多句念你」")
    body(
        doc,
        "用户投稿「想对某个人说的一句话」，精选100/1000句，形成长图、纪录内容、公众号专题——Campaign结束后留下内容资产。",
    )
    h2(doc, "传播节奏")
    table(
        doc,
        ["阶段", "重点"],
        [
            ["预热", "侨批文化短内容 + Hero Idea预告"],
            ["爆发", "产品上市 +「送给TA」教程 + KOL轻表达"],
            ["长尾", "许多句念你精选 + 用户故事专题"],
        ],
    )
    body(doc, "UGC以「被看见」激励为主；侨批翻译包装为文化体验，避免「AI噱头」压过情绪。")

    # ========== 12 资源预算 ==========
    h1(doc, "12｜项目资源配置（示意）与小规模验证")
    body(
        doc,
        "首阶段采用「小规模验证」，避免全国推广前产生供应链和执行风险。——资源有限时，先证明销售、会员与内容三端闭环。",
    )
    h2(doc, "第一阶段：区域试点")
    bullet(doc, "范围：华南核心城市 + 上海/北京内容城市（文化接近 + 传播扩散）。")
    table(
        doc,
        ["资源端", "投入项（示意）", "对应产出"],
        [
            ["产品端", "新品研发、包材设计、杯套生产", "可售SKU、货架识别、销量"],
            ["门店端", "门店物料、信箱装置、员工一句培训", "停留、参与、核销可执行"],
            ["内容端", "主视觉、短视频、KOL合作", "曝光、UGC、行动转化"],
            ["IP端", "授权费用、内容审核", "合规使用与信任背书"],
            ["数字端", "送给TA链路、电子侨批、会员埋点", "会员资产与二次触达"],
        ],
    )
    note(
        doc,
        "预算意识说明：本方案不虚构总金额。正式立项应按「区域试点框架」分项报价（研发/包材/物料/媒介/授权/技术），全国推广预算单列，以Phase 1指标门禁决定是否加码。",
    )

    # ========== 13 KPI ==========
    h1(doc, "13｜效果评估指标（KPI体系）")
    body(doc, "不设虚假精确数字；先建立可复盘指标体系，试点后用POS与后台回填目标值。")
    table(
        doc,
        ["类别", "指标", "看什么"],
        [
            ["销售", "联名销量、套餐购买比例、礼盒/电子礼销售、客单价", "是否卖得出、客单是否升"],
            ["用户", "信箱参与人数、电子侨批/纪念卡生成量、新会员增长、权益核销/二次打开", "是否沉淀用户资产"],
            ["内容", "UGC数量、话题互动量、用户故事数量、分享率", "是否形成传播记忆点"],
            ["品牌", "用户情感认知提升（有温度/会念你等，前后测或社媒语义）", "是否留下品牌资产"],
        ],
    )
    h2(doc, "分阶段门禁")
    bullet(doc, "Phase 1 试点验证：销量接受度、核销率、参与率、会员绑定率、店员可执行性。")
    bullet(doc, "Phase 2 内容扩散：UGC、话题、电子侨批分享率是否达扩散门槛。")
    bullet(doc, "Phase 3 资产沉淀：精选句子库、专题完整度、节点复用可行性。")

    # ========== 14 执行 ==========
    h1(doc, "14｜执行计划")
    table(
        doc,
        ["阶段", "目标", "关键动作", "关键指标"],
        [
            ["Phase 1", "小规模验证", "核心款、「送给TA」最小闭环、信箱、会员权益", "销量/核销/绑定/店员反馈"],
            ["Phase 2", "内容扩散", "Hero主推、KOL、许多句念你开启", "UGC/分享/话题"],
            ["Phase 3", "资产沉淀", "精选专题、复盘方法论、评估节日二波与扩城", "内容资产/复用条件"],
        ],
    )
    body(
        doc,
        "排期：以授权与供应链为准；中秋/重阳/春节/母亲节可作为第二波节点，而非唯一窗口。",
    )

    # ========== 15 风险 ==========
    h1(doc, "15｜风险与授权")
    table(
        doc,
        ["类型", "风险", "对策"],
        [
            ["商业", "IP热度下降", "不依赖单一档期；「念你」沉淀为节点资产"],
            ["商业", "购买动力不足", "口味价值+礼赠/异地场景+双杯"],
            ["商业", "活动参与成本高", "轻互动；嵌入既有点单/外卖路径"],
            ["商业", "只曝光不沉淀", "电子侨批强制轻会员路径，纳入KPI"],
            ["执行", "门店过重/核销乱", "一句口播；票根先试点"],
            ["合规", "版权金句误用", "授权清单+法务；未授权不宣称官联"],
            ["舆情", "情绪变沉重/消费苦难", "轻声惦记；聚焦念你与报平安"],
            ["产品", "供应链难撑复杂创新", "风味简单可量产；备胎风味"],
        ],
    )
    bullet(doc, "授权清单：片名、主视觉、形象、金句范围与期限；包装/杯套/短笺/电子侨批/专题审核。")
    bullet(doc, "数据复核：票房与联名先例上线前再核最新公开信息。")

    # Appendix mockups
    h1(doc, "附录A｜商业Mockup")
    body(doc, "三张决策用图：概念能成立、产品能卖、门店能执行。", False)
    h2(doc, "A1｜KV：给想念的人，一杯茶 / 心中念你")
    try_add_image(
        doc,
        MOCKUP_DIR / "01-kv-poster.png",
        caption="古典标题 + 繁体侨批手书金句 + 《给阿嬷的情书》电影标识，确保一眼读出联名。",
    )
    h2(doc, "A2｜产品包装（保留）")
    try_add_image(
        doc,
        MOCKUP_DIR / "02-product-packaging.png",
        caption="「念你」杯套 + 家书卡 + 明信片，证明货架可售。",
    )
    h2(doc, "A3｜门店体验")
    try_add_image(
        doc,
        MOCKUP_DIR / "03-store-experience.png",
        caption="念你信箱 + 侨批/海船/木棉电影主题墙，证明可执行。",
    )

    h1(doc, "附录B｜汇报目录（约15页）")
    table(
        doc,
        ["页", "标题"],
        [
            ["01", "项目判断：为什么值得做"],
            ["02", "项目目标：商业/品牌/用户"],
            ["03", "用户洞察"],
            ["04", "市场与品牌机会"],
            ["05", "IP资产矩阵"],
            ["06", "合作价值"],
            ["07", "Campaign架构与Hero Idea"],
            ["08", "给想念的人，一杯茶"],
            ["09", "产品与销售场景"],
            ["10", "念你信箱与会员"],
            ["11", "内容传播与资产"],
            ["12", "资源配置（小规模验证）"],
            ["13", "KPI体系"],
            ["14", "执行计划"],
            ["15", "风险与授权"],
        ],
    )

    h1(doc, "结语")
    quote(doc, "给想念的人，一杯茶。｜心中念你，便不觉远。")
    body(
        doc,
        "本案要回答三句：为什么值得投入；做完得到什么销售、会员与品牌结果；如何用小规模验证控制资源风险。主题不换，把商业落地写清楚。",
    )
    body(doc, "——完——", False)

    doc.save(str(OUT))
    print("saved", OUT)


if __name__ == "__main__":
    build()
